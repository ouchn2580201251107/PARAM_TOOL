"""
任务书视图模块
提供任务书的增删改查和导出功能
"""
import logging
from django.shortcuts import render, redirect
from django.views import View
from django.utils.decorators import method_decorator
from django.http import HttpResponse
import docx

from ..auth_views import login_required, role_required
from ..models import TaskDocument, Requirement

logger = logging.getLogger(__name__)


@method_decorator(login_required, name='dispatch')
class TaskDocumentListView(View):
    """
    任务书列表视图
    展示所有任务书文档，支持按需求编号和文档类型筛选
    """
    def get(self, request):
        logger.info(f"[TaskDocumentListView] 开始查询任务书列表，请求路径: {request.path}")
        try:
            docs = TaskDocument.objects.all().order_by('-generated_at')
            logger.info(f"[TaskDocumentListView] 任务书查询完成，记录数: {docs.count()}")
            
            req_no_filter = request.GET.get('req_no', '')
            type_filter = request.GET.get('type', '')
            logger.info(f"[TaskDocumentListView] 查询条件 - req_no: '{req_no_filter}', type: '{type_filter}'")
            
            if req_no_filter:
                docs = docs.filter(requirement__requirement_no__icontains=req_no_filter)
                logger.info(f"[TaskDocumentListView] 按需求编号过滤后记录数: {docs.count()}")
            
            if type_filter:
                docs = docs.filter(document_type=type_filter)
                logger.info(f"[TaskDocumentListView] 按文档类型过滤后记录数: {docs.count()}")
            
            is_editable = request.user.role.role_code in ['admin', 'technical']
            
            context = {
                'docs': docs,
                'req_no_filter': req_no_filter,
                'type_filter': type_filter,
                'types': TaskDocument.DOCUMENT_TYPE_CHOICES,
                'is_editable': is_editable,
            }
            logger.info(f"[TaskDocumentListView] 任务书列表查询成功，即将渲染模板")
            return render(request, 'parameter/task_document_list.html', context)
        except Exception as e:
            logger.error(f"[TaskDocumentListView] 任务书列表查询失败: {str(e)}", exc_info=True)
            raise


@method_decorator(role_required(['admin', 'technical']), name='dispatch')
class TaskDocumentCreateView(View):
    """
    任务书创建视图
    仅管理员和技术人员可访问，用于创建新的任务书
    """
    def get(self, request):
        logger.info(f"[TaskDocumentCreateView] 加载创建任务书页面")
        requirements = Requirement.objects.all()
        context = {
            'requirements': requirements,
            'types': TaskDocument.DOCUMENT_TYPE_CHOICES,
        }
        return render(request, 'parameter/task_document_create.html', context)
    
    def post(self, request):
        logger.info(f"[TaskDocumentCreateView] 创建任务书")
        try:
            task_doc = TaskDocument(
                requirement_id=request.POST.get('requirement'),
                document_no=request.POST.get('document_no'),
                title=request.POST.get('title'),
                document_type=request.POST.get('document_type'),
                content=request.POST.get('content'),
            )
            task_doc.save()
            logger.info(f"[TaskDocumentCreateView] 任务书创建成功，document_no: {task_doc.document_no}")
            return redirect('task_document_list')
        except Exception as e:
            logger.error(f"[TaskDocumentCreateView] 创建任务书失败: {str(e)}", exc_info=True)
            requirements = Requirement.objects.all()
            context = {
                'requirements': requirements,
                'types': TaskDocument.DOCUMENT_TYPE_CHOICES,
                'error': str(e),
                'data': request.POST,
            }
            return render(request, 'parameter/task_document_create.html', context)


@method_decorator(role_required(['admin', 'technical']), name='dispatch')
class TaskDocumentEditView(View):
    """
    任务书编辑视图
    仅管理员和技术人员可访问，用于修改任务书信息
    """
    def get(self, request, doc_id):
        logger.info(f"[TaskDocumentEditView] 加载编辑任务书页面，doc_id: {doc_id}")
        try:
            doc = TaskDocument.objects.get(id=doc_id)
            requirements = Requirement.objects.all()
            context = {
                'doc': doc,
                'requirements': requirements,
                'types': TaskDocument.DOCUMENT_TYPE_CHOICES,
            }
            return render(request, 'parameter/task_document_edit.html', context)
        except TaskDocument.DoesNotExist:
            logger.error(f"[TaskDocumentEditView] 任务书不存在，doc_id: {doc_id}")
            return render(request, 'parameter/error.html', {'message': '任务书不存在'})
    
    def post(self, request, doc_id):
        logger.info(f"[TaskDocumentEditView] 更新任务书，doc_id: {doc_id}")
        try:
            doc = TaskDocument.objects.get(id=doc_id)
            doc.requirement_id = request.POST.get('requirement')
            doc.document_no = request.POST.get('document_no')
            doc.title = request.POST.get('title')
            doc.document_type = request.POST.get('document_type')
            doc.content = request.POST.get('content')
            doc.save()
            logger.info(f"[TaskDocumentEditView] 任务书更新成功，document_no: {doc.document_no}")
            return redirect('task_document_list')
        except TaskDocument.DoesNotExist:
            logger.error(f"[TaskDocumentEditView] 任务书不存在，doc_id: {doc_id}")
            return render(request, 'parameter/error.html', {'message': '任务书不存在'})
        except Exception as e:
            logger.error(f"[TaskDocumentEditView] 更新任务书失败: {str(e)}", exc_info=True)
            doc = TaskDocument.objects.get(id=doc_id)
            requirements = Requirement.objects.all()
            context = {
                'doc': doc,
                'requirements': requirements,
                'types': TaskDocument.DOCUMENT_TYPE_CHOICES,
                'error': str(e),
            }
            return render(request, 'parameter/task_document_edit.html', context)


@method_decorator(role_required(['admin', 'technical']), name='dispatch')
class TaskDocumentDeleteView(View):
    """
    任务书删除视图
    仅管理员和技术人员可访问，用于删除任务书
    """
    def get(self, request, doc_id):
        logger.info(f"[TaskDocumentDeleteView] 删除任务书，doc_id: {doc_id}")
        try:
            doc = TaskDocument.objects.get(id=doc_id)
            doc.delete()
            logger.info(f"[TaskDocumentDeleteView] 任务书删除成功，document_no: {doc.document_no}")
            return redirect('task_document_list')
        except TaskDocument.DoesNotExist:
            logger.error(f"[TaskDocumentDeleteView] 任务书不存在，doc_id: {doc_id}")
            return render(request, 'parameter/error.html', {'message': '任务书不存在'})


@method_decorator(login_required, name='dispatch')
class TaskDocumentExportView(View):
    """
    任务书导出视图
    将任务书内容导出为Word文档
    """
    def get(self, request, doc_id):
        logger.info(f"[TaskDocumentExportView] 导出任务书，doc_id: {doc_id}")
        try:
            doc = TaskDocument.objects.get(id=doc_id)
            logger.info(f"[TaskDocumentExportView] 任务书查询成功，document_no: {doc.document_no}")
            
            document = docx.Document()
            document.add_heading(doc.title, level=1)
            document.add_paragraph(doc.content)
            
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = f'attachment; filename="{doc.document_no}.docx"'
            document.save(response)
            
            doc.exported_at = doc.generated_at
            doc.save()
            logger.info(f"[TaskDocumentExportView] 任务书导出成功，document_no: {doc.document_no}")
            
            return response
        except TaskDocument.DoesNotExist:
            logger.error(f"[TaskDocumentExportView] 任务书不存在，doc_id: {doc_id}")
            return render(request, 'parameter/error.html', {'message': '任务书不存在'})
        except Exception as e:
            logger.error(f"[TaskDocumentExportView] 任务书导出失败，doc_id: {doc_id}, 错误: {str(e)}", exc_info=True)
            raise
